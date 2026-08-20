import os
import hashlib
from datetime import datetime
from pathlib import Path
from pypdf import PdfReader
import pandas as pd
import docx
from langchain_core.documents import Document

class DocumentLoader:
    def __init__(self):
        pass

    def load_file(self, file_path: str) -> Document:
        """Loads a single file (PDF, TXT, MD, DOCX, XLSX, CSV) and returns a Document with metadata."""
        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext in (".txt", ".md"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        elif ext == ".pdf":
            reader = PdfReader(path)
            pages_content = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_content.append(text)
            content = "\n".join(pages_content)
        elif ext == ".docx":
            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract text from tables in docx as well
            table_lines = []
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_lines.append(" | ".join(row_cells))
            
            content = "\n".join(paragraphs)
            if table_lines:
                content += "\n\n--- Document Tables ---\n" + "\n".join(table_lines)
        elif ext == ".csv":
            df = pd.read_csv(path)
            content = f"CSV File: {path.name}\nColumns: {', '.join(df.columns)}\n\nData Rows:\n" + df.to_string(index=False)
        elif ext in (".xlsx", ".xls"):
            excel = pd.ExcelFile(path)
            sheet_contents = []
            for sheet_name in excel.sheet_names:
                df = excel.parse(sheet_name)
                sheet_text = f"--- Sheet: {sheet_name} ---\nColumns: {', '.join([str(c) for c in df.columns])}\n\nData Rows:\n" + df.to_string(index=False)
                sheet_contents.append(sheet_text)
            content = f"Excel Workbook: {path.name}\n\n" + "\n\n".join(sheet_contents)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        # Compute MD5 hash of raw content
        md5_hash = hashlib.md5(content.encode("utf-8")).hexdigest()

        # Extract file times
        stat = path.stat()
        created_time = datetime.fromtimestamp(stat.st_ctime).isoformat()
        modified_time = datetime.fromtimestamp(stat.st_mtime).isoformat()

        metadata = {
            "source": path.name,
            "file_path": str(path),
            "created_at": created_time,
            "modified_at": modified_time,
            "doc_id": md5_hash
        }

        return Document(page_content=content, metadata=metadata)

    def load_directory(self, dir_path: str) -> list[Document]:
        """Recursively loads all supported files (.pdf, .txt, .md, .docx, .xlsx, .csv) from directory."""
        documents = []
        path = Path(dir_path).resolve()
        
        if not path.is_dir():
            raise ValueError(f"Not a directory: {dir_path}")

        supported_extensions = {".pdf", ".txt", ".md", ".docx", ".xlsx", ".xls", ".csv"}
        
        for file in path.rglob("*"):
            if file.is_file() and file.suffix.lower() in supported_extensions:
                try:
                    doc = self.load_file(str(file))
                    documents.append(doc)
                except Exception as e:
                    print(f"Error loading {file}: {e}")
                    
        return documents
